#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os
import glob
import re
import json
import time
from typing import Any, Dict, List, Optional
from fastapi import FastAPI, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel

current_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, current_dir)
sys.path.insert(0, os.path.join(current_dir, 'src'))

outer_env_path = os.path.join(os.path.dirname(current_dir), '.env')
from dotenv import load_dotenv
load_dotenv(outer_env_path, override=True)
load_dotenv(os.path.join(current_dir, '.env'), override=False)

from core import settings
from core import logger
from agents import agent_factory, FallbackAgent
from tools import discover_tools, ToolRegistry, get_all_tools

from langchain_core.messages import HumanMessage, AIMessage, BaseMessage, SystemMessage

app = FastAPI(title="X-LangChain API", description="LangChain 智能助手 API", version="1.0.0")

# ── RAG 文档上传目录 ──
import shutil
RAG_DOCS_DIR = os.path.join(current_dir, "data", "rag", "docs")
os.makedirs(RAG_DOCS_DIR, exist_ok=True)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

agent = None

# 会话记忆存储: { session_id: [HumanMessage, AIMessage, ...] }
_session_store: Dict[str, List[BaseMessage]] = {}
MAX_MEMORY_MESSAGES = 200  # 每个会话最多保留的消息数（压缩后仍保留完整历史）

# 上下文压缩器
from core.context_compressor import get_compressor, estimate_tokens
_compressor = get_compressor()

# 用户自定义 System Prompt（通过设置面板配置）
_custom_system_prompt: str = ""

# 语义会话记忆
from core.semantic_memory import get_semantic_memory
_semantic_memory = get_semantic_memory()

# 生成式 Agent (记忆流 + 反思 + 规划)
from agents.generative_agent import get_generative_agent
_gen_agent = get_generative_agent()


def _get_history(session_id: str) -> List[BaseMessage]:
    """获取会话历史消息列表。"""
    if session_id not in _session_store:
        _session_store[session_id] = []
    return _session_store[session_id]


def _build_file_context(files: list, max_chars_per_file: int = 8000) -> str:
    """构建文件上下文文本，追加到用户消息中。"""
    if not files:
        return ""
    parts = ["\n\n--- 用户上传文件 ---"]
    for f in files:
        name = f.get("name", "unknown")
        content = f.get("content", "")
        ft = f.get("type", "")
        # 截断太长文件
        if len(content) > max_chars_per_file:
            content = content[:max_chars_per_file] + "\n...(内容已截断)"
        parts.append(f"\n### 文件: {name}")
        if ft:
            parts.append(f"类型: {ft}")
        parts.append(f"内容:\n{content}")
    parts.append("--- 文件结束 ---\n")
    return "\n".join(parts)


def _get_compressed_context(session_id: str, current_msg: str) -> List[BaseMessage]:
    """
    获取压缩 + 语义记忆 + 生成式记忆增强后的上下文。

    流程：
    1. 生成式记忆流检索（重要性+时效+相关性加权）
    2. 语义检索相关历史记忆（跨会话）
    3. 取完整历史 → 追加当前消息
    4. 超过阈值时触发压缩（摘要 + 滑动窗口）
    5. 语义记忆 / 生成记忆 / 反思洞察 / 活跃计划 作为 SystemMessage 注入
    6. 返回完整上下文发给 Agent（Agent 自带 System Prompt）

    最终结构：
      [SystemMessage(生成式记忆)] + [SystemMessage(语义记忆)] + [SystemMessage(历史摘要)] + 最近 N 条消息
    """
    history = _get_history(session_id)
    all_messages = history + [HumanMessage(content=current_msg)]

    compressed = _compressor.get_compressed_context(session_id, all_messages)

    # 记录压缩前后令牌对比
    raw_tokens = estimate_tokens(
        " ".join(m.content for m in all_messages if hasattr(m, "content"))
    )
    comp_tokens = estimate_tokens(
        " ".join(m.content for m in compressed if hasattr(m, "content"))
    )
    if len(compressed) < len(all_messages):
        logger.info(
            f"上下文压缩生效 [{session_id[:20]}]: "
            f"{len(all_messages)}条 ({raw_tokens} tokens) → "
            f"{len(compressed)}条 ({comp_tokens} tokens)"
        )

    # ── 生成式记忆注入 (记忆流 + 反思 + 计划) ──
    try:
        gen_context = _gen_agent.build_context(current_msg, session_id)
        if gen_context:
            compressed.insert(0, SystemMessage(content=gen_context))
    except Exception as e:
        logger.warning(f"生成式记忆注入失败: {e}")

    # ── 语义记忆注入 ──
    try:
        memories = _semantic_memory.retrieve(current_msg)
        if memories:
            memory_text = _semantic_memory.format_for_context(memories)
            compressed.insert(0, SystemMessage(content=memory_text))
            logger.info(f"语义记忆注入: 从 {len(_semantic_memory._memory_cache)} 条中召回 {len(memories)} 条")
    except Exception as e:
        logger.warning(f"语义记忆检索失败: {e}")

    # ── 用户自定义 System Prompt 注入（设置的文案追加在 Agent 自带 Prompt 之前） ──
    if _custom_system_prompt:
        compressed.insert(0, SystemMessage(content=_custom_system_prompt))

    return compressed


def _add_to_history(session_id: str, user_msg: str, ai_msg: str) -> None:
    """向会话历史添加一轮对话，超出 MAX_MEMORY_MESSAGES 时裁剪旧消息。"""
    msgs = _get_history(session_id)
    msgs.append(HumanMessage(content=user_msg))
    msgs.append(AIMessage(content=ai_msg))
    while len(msgs) > MAX_MEMORY_MESSAGES:
        msgs.pop(0)
        msgs.pop(0)


def _tool_label(tool_name: str) -> str:
    """工具名 → 中文标签。"""
    mapping = {
        "web_search": "联网搜索",
        "web_search_tool": "联网搜索",
        "weather_search_tool": "天气查询",
        "exchange_rate_tool": "汇率查询",
        "search_calendar": "日历查询",
        "get_current_datetime": "获取时间",
        "code_sandbox_tool": "代码执行",
        "rag_search_tool": "知识库检索",
    }
    return mapping.get(tool_name, tool_name)


def _extract_search_sources(response_text: str, tools_used: list) -> list:
    """提取搜索来源。优先从 web_tool 侧通道获取，兜底从回复文本扫描 URL。"""
    SEARCH_TOOLS = {"web_search", "web_search_tool", "duckduckgo_search", "tavily_search"}
    if not any(t.lower() in {s.lower() for s in SEARCH_TOOLS} for t in tools_used):
        return []

    # ── 优先：从 web_tool 侧通道获取（精确的标题+URL+摘要）──
    try:
        from tools.web_tool import get_last_search_meta
        meta = get_last_search_meta()
        if meta:
            return [{"title": m.get("title", "")[:80],
                     "url": m.get("url", ""),
                     "snippet": m.get("snippet", "")[:200]}
                    for m in meta if m.get("url") or m.get("title")]
    except Exception:
        pass

    # ── 兜底：从回复文本扫描 URL ──
    sources = []
    seen = set()
    url_pattern = re.compile(r'https?://[^\s<>"\')\]}，。；;]+')
    for match in url_pattern.finditer(response_text):
        url = match.group(0).rstrip('.')
        if url in seen:
            continue
        seen.add(url)
        line_start = max(0, match.start() - 200)
        context = response_text[line_start:match.end()]
        md_link = re.search(r'\[([^\]]*)\]\(' + re.escape(url) + r'\)', context)
        title = md_link.group(1) if md_link else url
        sources.append({"title": title[:80], "url": url, "snippet": ""})

    for md_match in re.finditer(r'\[([^\]]+)\]\((https?://[^\)]+)\)', response_text):
        title, url = md_match.group(1), md_match.group(2).rstrip('.')
        if url not in seen:
            seen.add(url)
            sources.append({"title": title[:80], "url": url, "snippet": ""})

    return sources[:8]

class ChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = None
    model_name: Optional[str] = None
    files: Optional[List[dict]] = None  # [{name, type, content}, ...]

class ChatResponse(BaseModel):
    response: str
    model: str

@app.on_event("startup")
async def startup_event():
    global agent
    try:
        # 自动发现并注册所有 Skill/Tool
        logger.info("正在发现并注册所有 Skill/Tool...")
        discovered = discover_tools()
        stats = ToolRegistry.get_stats()
        logger.info(
            f"Skill 加载完成: 总计 {stats['total_tools']} 个, "
            f"分类 {stats['categories']} 个"
        )
        for tool_name in ToolRegistry.get_names():
            metadata = ToolRegistry.get_metadata(tool_name) or {}
            logger.info(
                f"  - [Skill] {tool_name}: {metadata.get('description', '无描述')}"
            )

        model_name = os.getenv("MODEL_NAME", "mock")
        logger.info(f"正在初始化 Agent (使用 {model_name} 模型)...")
        if not settings.validate_model_config(model_name):
            raise ValueError(f"{model_name} 模型配置不完整")
        agent = FallbackAgent(primary_model=model_name)
        logger.info("Agent 初始化成功（已启用会话记忆 + 链路回退）")

        # 预加载 Embedding 模型，避免首次请求时阻塞等待下载
        try:
            from rag.vector_store import _get_embeddings
            logger.info("正在预加载 Embedding 模型...")
            _get_embeddings()
            logger.info("Embedding 模型预加载完成")
        except Exception as emb_e:
            logger.warning(f"Embedding 模型预加载失败: {emb_e}")
    except Exception as e:
        logger.error(f"Agent 初始化失败: {e}")
        raise

@app.get("/health")
async def health_check():
    return {"status": "healthy", "model": os.getenv("MODEL_NAME", "mock"), "fallback": True}

@app.get("/skills")
async def list_skills():
    """返回所有已注册的 Skill/Tool 列表"""
    skills = []
    for tool_name in ToolRegistry.get_names():
        metadata = ToolRegistry.get_metadata(tool_name) or {}
        skills.append({
            "name": tool_name,
            "category": metadata.get("category", "default"),
            "subcategory": metadata.get("subcategory"),
            "description": metadata.get("description", ""),
        })
    stats = ToolRegistry.get_stats()
    return {
        "total": stats["total_tools"],
        "categories": stats["categories"],
        "skills": skills,
    }

@app.post("/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    if not agent:
        raise HTTPException(status_code=500, detail="Agent 未初始化")
    
    try:
        logger.info(f"收到请求: {request.message[:50]}...")
        
        session_id = request.session_id or "default"
        # 追加文件上下文到消息
        message = request.message + _build_file_context(request.files)
        # 使用压缩后的上下文
        all_messages = _get_compressed_context(session_id, message)
        logger.info(f"会话 {session_id}: 上下文共 {len(all_messages)} 条消息")

        response, model_used = await agent.invoke(all_messages)

        # 保存到历史
        _add_to_history(session_id, request.message, response)

        # 存入语义记忆
        try:
            _semantic_memory.add_memory(session_id, request.message, response)
        except Exception:
            pass

        # 生成式 Agent: 观察 + 尝试规划
        try:
            _gen_agent.observe(session_id, request.message, response)
            _gen_agent.try_plan(request.message, session_id)
        except Exception:
            pass

        return {"response": response, "model": model_used}
    
    except Exception as e:
        logger.error(f"处理请求时出错: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/chat/stream")
async def chat_stream(request: ChatRequest):
    if not agent:
        raise HTTPException(status_code=500, detail="Agent 未初始化")

    logger.info(f"收到流式请求: {request.message[:50]}...")

    from constants import STREAM_MODE_MESSAGES

    session_id = request.session_id or "default"
    # 追加文件上下文到消息
    message = request.message + _build_file_context(request.files)
    # 使用压缩后的上下文
    all_messages = _get_compressed_context(session_id, message)
    logger.info(f"会话 {session_id}: 上下文共 {len(all_messages)} 条消息（流式 + 链路回退）")

    async def _stream_generator():
        full_response = ""
        model_used = ""
        done = False
        stream_start = time.time()       # 流开始时间
        last_think_ts = 0                 # 上次思考节点发送时间
        first_token_emitted = False
        thinking_started = False
        # 无工具调用的纯推理流：阶梯式思考节点（字符阈值 + 时间间隔双重管控）
        think_progress = {"reason": False, "validate": False, "fix": False}
        tools_used = []
        sources_emitted = False

        try:
            async for token, meta, model in agent.astream(all_messages, stream_mode=STREAM_MODE_MESSAGES):
                if meta is not None:
                    if isinstance(meta, dict):
                        tc_list = meta.get("_tool_calls", [])
                        if tc_list:
                            if not thinking_started:
                                thinking_started = True
                                last_think_ts = time.time()
                                yield {"thinking": [{"type": "decompose", "text": "拆解问题"}]}
                                think_progress["reason"] = True
                                yield {"thinking": [{"type": "reason", "text": "推理分析"}]}

                            for tc_name in tc_list:
                                if tc_name not in tools_used:
                                    tools_used.append(tc_name)
                                    label = _tool_label(tc_name)
                                    yield {"thinking": [{"type": "tools", "text": f"调用: {label}"}]}

                        if meta.get("_tool_result"):
                            tool_name = meta.get("_tool_name", "")
                            if tool_name and tool_name not in tools_used:
                                tools_used.append(tool_name)
                            yield {"thinking": [{"type": "validate", "text": "校验返回结果"}]}
                            yield {"thinking": [{"type": "fix", "text": "修正优化"}]}
                            yield {"thinking": [{"type": "summary", "text": "总结输出"}]}
                            if tool_name.lower() in {"web_search", "web_search_tool"}:
                                try:
                                    from tools.web_tool import get_last_search_meta
                                    meta_list = get_last_search_meta()
                                    if meta_list:
                                        sources = [{"title": m.get("title","")[:80],
                                                    "url": m.get("url",""),
                                                    "snippet": m.get("snippet","")[:200]}
                                                   for m in meta_list if m.get("url") or m.get("title")]
                                        if sources:
                                            sources_emitted = True
                                            yield {"sources": sources}
                                except Exception:
                                    pass
                            continue

                    if token:
                        model_used = model or model_used
                        full_response += token
                        if not first_token_emitted:
                            first_token_emitted = True
                            thinking_started = True
                            last_think_ts = time.time()
                            yield {"thinking": [{"type": "decompose", "text": "拆解问题"}]}
                        # ── 无工具调用的纯推理流：字符位置 + 时间间隔双重管控 ──
                        # 每个思考节点至少间隔 1.8s，避免一股脑出现
                        if not tools_used:
                            now = time.time()
                            elapsed = now - last_think_ts
                            clen = len(full_response)
                            if not think_progress["reason"] and clen >= 200 and elapsed >= 1.8:
                                think_progress["reason"] = True
                                last_think_ts = now
                                yield {"thinking": [{"type": "reason", "text": "推理分析"}]}
                            elif not think_progress["validate"] and clen >= 600 and elapsed >= 1.8:
                                think_progress["validate"] = True
                                last_think_ts = now
                                yield {"thinking": [{"type": "validate", "text": "校验分析"}]}
                            elif not think_progress["fix"] and clen >= 1200 and elapsed >= 1.8:
                                think_progress["fix"] = True
                                last_think_ts = now
                                yield {"thinking": [{"type": "fix", "text": "归纳整合"}]}
                        yield {"token": token, "done": False}
                elif token and meta is None and not done:
                    model_used = model or model_used
                    full_response = token
                    yield {"token": token, "done": False}
                else:
                    done = True

            if not sources_emitted:
                sources = _extract_search_sources(full_response, tools_used)
                if sources:
                    yield {"sources": sources}

            if thinking_started:
                if not tools_used:
                    yield {"thinking": [{"type": "summary", "text": "总结输出"}]}
                yield {"thinking_end": True}

            # 保存到历史
            _add_to_history(session_id, request.message, full_response)

            # 存入语义记忆
            try:
                _semantic_memory.add_memory(session_id, request.message, full_response)
            except Exception:
                pass

            # 生成式 Agent: 观察 + 尝试规划
            try:
                _gen_agent.observe(session_id, request.message, full_response)
                _gen_agent.try_plan(request.message, session_id)
            except Exception:
                pass

            yield {"token": "", "done": True}

        except Exception as e:
            logger.error(f"流式请求处理出错: {e}")
            yield {"token": f"抱歉，处理请求时出错: {e}", "done": True}

    async def _jsonl_wrapper():
        async for item in _stream_generator():
            yield json.dumps(item, ensure_ascii=False) + "\n"

    return StreamingResponse(
        _jsonl_wrapper(),
        media_type="application/jsonl; charset=utf-8",
    )

@app.get("/logs")
async def query_logs(
    keyword: Optional[str] = Query(None, description="模糊搜索关键词"),
    level: Optional[str] = Query(None, description="日志级别: INFO, ERROR, DEBUG, WARNING"),
    limit: int = Query(100, description="返回条数限制"),
    offset: int = Query(0, description="起始偏移量")
):
    """查询日志（支持模糊搜索）"""
    log_dir = os.path.join(current_dir, 'logs')
    if not os.path.exists(log_dir):
        return {"total": 0, "logs": []}
    
    log_files = sorted(glob.glob(os.path.join(log_dir, '*.log')), reverse=True)
    all_logs = []
    
    for log_file in log_files:
        try:
            with open(log_file, 'r', encoding='utf-8', errors='ignore') as f:
                for line in f:
                    line = line.strip()
                    if not line:
                        continue
                    
                    if level and level.upper() not in line:
                        continue
                    
                    if keyword and keyword.lower() not in line.lower():
                        continue
                    
                    all_logs.append(line)
                    if len(all_logs) >= limit + offset:
                        break
        except Exception:
            continue
        if len(all_logs) >= limit + offset:
            break
    
    total = len(all_logs)
    result_logs = all_logs[offset:offset + limit]
    
    return {"total": total, "logs": result_logs}

@app.get("/config/model")
async def get_model_config():
    """获取当前模型配置"""
    return {
        "temperature": settings.TEMPERATURE,
        "model_name": settings.MODEL_NAME,
        "system_prompt": _custom_system_prompt,
        "openai_api_key": settings.OPENAI_API_KEY if settings.OPENAI_API_KEY else "",
        "openai_base_url": settings.OPENAI_BASE_URL if settings.OPENAI_BASE_URL else "",
        "mimo_model": settings.MIMO_MODEL if settings.MIMO_MODEL else "",
        "deepseek_api_key": settings.DEEPSEEK_API_KEY if settings.DEEPSEEK_API_KEY else "",
        "doubao_api_key": settings.DOUBAO_API_KEY if settings.DOUBAO_API_KEY else "",
        "aliyun_api_key": settings.ALIYUN_API_KEY if settings.ALIYUN_API_KEY else "",
        "debug": settings.DEBUG,
        "structured": settings.STRUCTURED,
    }

class ModelConfigRequest(BaseModel):
    temperature: Optional[float] = None
    model_name: Optional[str] = None
    max_tokens: Optional[int] = None
    system_prompt: Optional[str] = None

@app.post("/config/model")
async def update_model_config(config: ModelConfigRequest):
    """更新模型配置（温度、模型名称等）"""
    global _custom_system_prompt
    updates = {}
    
    if config.temperature is not None:
        if 0 <= config.temperature <= 2:
            os.environ['TEMPERATURE'] = str(config.temperature)
            settings.TEMPERATURE = config.temperature
            updates['temperature'] = config.temperature
        else:
            raise HTTPException(status_code=400, detail="温度参数必须在 0-2 之间")
    
    if config.model_name:
        os.environ['MODEL_NAME'] = config.model_name
        updates['model_name'] = config.model_name
    
    if config.max_tokens:
        os.environ['MAX_TOKENS'] = str(config.max_tokens)
        updates['max_tokens'] = config.max_tokens

    if config.system_prompt is not None:
        _custom_system_prompt = config.system_prompt
        updates['system_prompt'] = 'updated'

    logger.info(f"模型配置已更新: {updates}")
    return {"success": True, "updated": updates}

class ExportRequest(BaseModel):
    content: str
    filename: Optional[str] = "agent-export"

@app.post("/export/docx")
async def export_docx(request: ExportRequest):
    """将 Markdown/文本内容导出为 .docx 文件"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    content = request.content.strip()
    lines = content.split('\n')

    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph('')
            continue

        if stripped.startswith('# ') or stripped.startswith('## ') or stripped.startswith('### '):
            level = stripped.count('#', 0, 3)
            heading_text = stripped.lstrip('#').strip()
            p = doc.add_heading(heading_text, level=min(level, 3))
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif re.match(r'^\d+[.)]\s', stripped):
            doc.add_paragraph(re.sub(r'^\d+[.)]\s*', '', stripped), style='List Number')
        else:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = re.sub(r'[\\/*?:"<>|]', '', request.filename)
    return StreamingResponse(
        buf,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{safe_name}.docx"'}
    )

# =============================================================================
# RAG 文档管理 API
# =============================================================================

from fastapi import UploadFile, File as FastAPIFile

@app.post("/rag/upload")
async def rag_upload(file: UploadFile = FastAPIFile(...)):
    """上传文档到 RAG 知识库（支持 PDF/Word/Markdown/TXT）。"""
    if not file.filename:
        raise HTTPException(status_code=400, detail="文件名不能为空")

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in (".pdf", ".docx", ".md", ".txt"):
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")

    try:
        from rag import get_rag_store

        file_path = os.path.join(RAG_DOCS_DIR, file.filename)
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)

        store = get_rag_store()
        chunk_count = store.add_document(file_path)

        return {
            "success": True,
            "filename": file.filename,
            "chunks": chunk_count,
        }
    except Exception as e:
        logger.error(f"RAG 文档上传失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/rag/documents")
async def rag_list_documents():
    """列出已上传的文档。"""
    try:
        from rag import get_rag_store
        store = get_rag_store()
        docs = store.list_documents()
        return {"total": len(docs), "documents": docs}
    except Exception as e:
        logger.error(f"RAG 文档列表失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/rag/documents/{filename}")
async def rag_delete_document(filename: str):
    """从知识库删除文档。"""
    try:
        from rag import get_rag_store
        store = get_rag_store()
        ok = store.delete_document(filename)
        if not ok:
            raise HTTPException(status_code=404, detail="文档不存在或删除失败")
        return {"success": True, "filename": filename}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"RAG 文档删除失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/rag/search")
async def rag_search(
    query: str = Query(..., description="搜索查询"),
    k: int = Query(4, description="返回结果数"),
    strategy: str = Query("hybrid", description="检索策略: semantic, bm25, mmr, hybrid, all"),
    expand: bool = Query(True, description="启用多面查询扩展（提高召回）"),
    rerank: bool = Query(True, description="启用 Cross-Encoder 精排（提高精度）"),
):
    """搜索 RAG 知识库（支持多策略融合检索 + 查询扩展 + 精排）。"""
    try:
        from rag import get_rag_store
        store = get_rag_store()
        results = store.fused_search(
            query, k=k, strategy=strategy,
            expand_query=expand, cross_encode_rerank=rerank,
        )
        return {
            "query": query,
            "strategy": strategy,
            "expand": expand,
            "rerank": rerank,
            "results": [{"content": item["content"][:500], "source": item["source"], "score": round(score, 4)} for item, score in results]
        }
    except Exception as e:
        logger.error(f"RAG 搜索失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/")
async def serve_frontend():
    return FileResponse(os.path.join(current_dir, "frontend", "index.html"))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001, log_level="info")